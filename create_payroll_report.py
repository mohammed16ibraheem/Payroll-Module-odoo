# -*- coding: utf-8 -*-
"""
Creates a professional Word document for Payroll Module task submission.
Includes ALL screenshots (1-76, 78-100) and complete documentation.
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# All images in order (77.png missing)
IMAGE_LIST = list(range(1, 77)) + list(range(78, 101))

# Captions for key images
CAPTIONS = {
    1: 'Odoo Apps — Payroll module (highlighted)',
    2: 'Payroll navigation bar',
    3: 'Configuration menu — Work Entry Types',
    4: 'Payroll dashboard',
    5: 'Payroll settings',
    6: 'Payslips menu',
    7: 'Loading state',
    8: 'Employee Payslips list',
    9: 'Payslip list with RAM LAL',
    10: 'AAVEZ NAUSHAD DALVI — January 2026 payslip (Test Case 2 ✓)',
    11: 'Payslip heading "January 2026" (Fix 2 verified)',
    12: 'RAM LAL payslip in list',
    13: 'RAM LAL payslip — Paid Time Off in Worked Days (BEFORE)',
    14: 'Configuration dropdown — Work Entry Types',
    15: 'RAM LAL payslip — Bug #1 and Bug #2 visible',
    16: 'Work Entry Types list — Paid Time Off (highlighted)',
    17: 'Paid Time Off — Time Off CHECKED (BEFORE)',
    18: 'Paid Time Off — Time Off CHECKED (BEFORE)',
    19: 'Paid Time Off — Time Off UNCHECKED (AFTER Fix 1 ✓)',
    20: 'Paid Time Off — Time Off UNCHECKED (AFTER Fix 1 ✓)',
    21: 'Paid Time Off — Time Off UNCHECKED (AFTER Fix 1 ✓)',
    22: 'Generic Time Off configuration',
    23: 'Generic Time Off — Time Off option',
    24: 'Generic Time Off — Time Off UNCHECKED',
    25: 'Generic Time Off configuration',
    26: 'Compensatory Time Off — Time Off CHECKED (BEFORE)',
    27: 'Compensatory Time Off — Time Off UNCHECKED (AFTER)',
    28: 'Compensatory Time Off — Time Off UNCHECKED (AFTER)',
    29: 'Compensatory Time Off — Time Off UNCHECKED (AFTER)',
    30: 'Unpaid work entry type',
    50: 'Settings → Technical → Views — report_payslip',
    60: 'Report template — Line 4 (use date_to for End Date)',
    70: 'Payslip confirmation dialog',
    80: 'RAM LAL payslip — Worked Days & Inputs',
    90: 'Employee Payslips list',
    100: 'Paid Time Off — Time Off UNCHECKED (Final state ✓)',
}

def add_image(doc, num, width=6.0):
    f = f'{num}.png'
    path = os.path.join(BASE_DIR, f)
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width))
            cap = CAPTIONS.get(num, f'Screenshot {num}')
            p = doc.add_paragraph(f'Figure {num}: {cap}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
            return True
        except Exception as e:
            doc.add_paragraph(f'[Image {f} error: {e}]')
    else:
        doc.add_paragraph(f'[Image {f} not found]')
    return False

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    # ===== TITLE =====
    title = doc.add_heading('Payroll Module — Task Completion Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('Submitted by: Namma Al Enjaz')
    doc.add_paragraph('Date: March 2026')
    doc.add_paragraph('Environment: erpnamma-stage-26999149.dev.odoo.com (Odoo)')
    doc.add_page_break()

    # ===== 1. TASK & BEFORE STATE =====
    doc.add_heading('1. Task Overview', 1)
    doc.add_paragraph('Both Tasks Completed & Tested Successfully', style='Intense Quote')
    doc.add_paragraph()

    doc.add_heading('1.1 BEFORE — What Was Found', 2)
    doc.add_paragraph('Bug #1 — Time Off Auto-Importing into Payslip:', style='Heading 3')
    doc.add_paragraph(
        'RAM LAL\'s payslip (Period: 21/11/2025 - 20/12/2025) had in its Worked Days tab: '
        'Paid Time Off → 12 days, 96 hours, 476.80 SR — automatically pulled from Time Off module, reducing salary.'
    )
    doc.add_paragraph('Reference: erpnamma-stage-26999149.dev.odoo.com (Payslip form)', style='List Bullet')
    doc.add_paragraph()

    doc.add_paragraph('Bug #2 — Wrong Month Heading:', style='Heading 3')
    doc.add_paragraph(
        'The same payslip had period ending December but the title showed "November 2025 Payslip" '
        '(using Start Date month instead of End Date month).'
    )
    doc.add_paragraph()

    # ===== 2. FIX 1 =====
    doc.add_heading('2. FIX 1 — Disable Time Off → Payroll Link', 1)
    doc.add_paragraph('Where to go: Payroll → Configuration → Work Entry Types', style='Intense Quote')
    doc.add_paragraph(
        'What was changed: Opened each leave-type work entry and unchecked the "Time Off" checkbox under the TIME OFF OPTIONS section.'
    )
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Work Entry Type'
    hdr[1].text = 'Code'
    hdr[2].text = 'Change'
    data = [
        ('Generic Time Off', 'LEAVE100', 'Time Off unchecked'),
        ('Compensatory Time Off', 'LEAVE105', 'Time Off unchecked'),
        ('Unpaid', 'LEAVE90', 'Time Off unchecked'),
        ('Sick Time Off', 'LEAVE110', 'Time Off unchecked'),
        ('Paid Time Off', 'LEAVE120', 'Time Off unchecked'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        row[2].text = row_data[2]
    
    doc.add_paragraph()
    doc.add_paragraph(
        'AFTER result: The "Time Off" checkbox is now UNCHECKED (empty) and the "Time Off Type" field is completely gone — '
        'validated leave approvals will no longer auto-import into payslip Worked Days. Manual salary adjustments are now the only way vacation affects salary.'
    )
    doc.add_paragraph()

    # ===== 3. FIX 2 =====
    doc.add_heading('3. FIX 2 — Payslip Report Month Heading → Use End Date', 1)
    doc.add_paragraph(
        'Where to go: Settings → Technical → User Interface → Views → Search report_payslip → Open hr_payroll.report_payslip',
        style='Intense Quote'
    )
    doc.add_paragraph('What was changed — Line 4 of the template:', style='Heading 3')
    doc.add_paragraph('BEFORE: (Used o.name which is computed from date_from / Start Date)', style='Heading 4')
    p = doc.add_paragraph()
    p.add_run('<h2 id="payslip_name"><span t-field="o.name">August 2023 Payslip</span></h2>').font.name = 'Consolas'
    p.add_run('').font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('AFTER: (Now uses o.date_to / End Date directly)', style='Heading 4')
    p2 = doc.add_paragraph()
    p2.add_run('<h2 id="payslip_name"><span t-esc="o.date_to.strftime(\'%B %Y\') + \' Payslip\'"/></h2>').font.name = 'Consolas'
    doc.add_paragraph()

    # ===== 4. TEST RESULTS =====
    doc.add_heading('4. TEST RESULTS', 1)
    doc.add_paragraph('Test Case 1 — RAM LAL (Period: 21/11/2025 → 20/12/2025):', style='Heading 3')
    doc.add_paragraph('BEFORE heading: "November 2025 Payslip"', style='List Bullet')
    doc.add_paragraph('AFTER heading: "December 2025 Payslip"', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Test Case 2 — AAVEZ NAUSHAD DALVI (Period: 01/01/2026 → 31/01/2026):', style='Heading 3')
    doc.add_paragraph('Heading: "January 2026 Payslip" (Same month — correct)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Test Case 3 — Paid Time Off Work Entry Type:', style='Heading 3')
    doc.add_paragraph('Time Off checkbox: UNCHECKED — link to Time Off module disabled', style='List Bullet')
    doc.add_paragraph()

    # ===== 5. IMPORTANT NOTES =====
    doc.add_heading('5. Important Notes', 1)
    doc.add_paragraph(
        '1. Existing payslips that were already computed will still show old Paid Time Off data — this is historical and unchanged. '
        'Only newly computed/recomputed payslips will no longer auto-import leave days.',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Future module updates may overwrite the QWeb report change (Odoo showed a warning). If this happens, you will need to re-apply Line 4 or use an Inherited View via Odoo Studio for a permanent solution.',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Alternative solution to investigate later: Instead of disabling the Time Off link entirely, you could keep it enabled but set the leave Work Entry Types to have zero amount/coefficient in the Salary Structure rules — this way leaves are still tracked in worked days for reporting but do not reduce salary.',
        style='List Number'
    )
    doc.add_paragraph()

    # ===== 6. REFERENCE LINKS =====
    doc.add_heading('6. Reference Links', 1)
    doc.add_paragraph('Payslip form: https://erpnamma-stage-26999149.dev.odoo.com (hr.payslip)')
    doc.add_paragraph('Work Entry Types: https://erpnamma-stage-26999149.dev.odoo.com (hr.work.entry.type)')
    doc.add_paragraph('Payslip Report RAM LAL: .../report/html/hr_payroll.report_payslip_lang/13129')
    doc.add_paragraph('Payslip Report AAVEZ: .../report/html/hr_payroll.report_payslip_lang/13433')
    doc.add_page_break()

    # ===== 7. SCREENSHOTS (ALL IMAGES) =====
    doc.add_heading('7. Screenshots & Evidence', 1)

    # ===== 3. SCREENSHOTS (ALL IMAGES) =====
    doc.add_heading('3. Screenshots & Evidence', 1)
    doc.add_paragraph('The following screenshots document the navigation, before/after states, and test results. All images are included in order.')

    added = 0
    for num in IMAGE_LIST:
        add_image(doc, num)
        doc.add_paragraph()
        added += 1

    doc.add_heading('8. Conclusion', 1)
    doc.add_paragraph(
        'Both tasks were completed and tested successfully. Time Off no longer auto-imports into payslips, and the payslip report heading correctly uses the period End Date month.'
    )

    out_path = os.path.join(BASE_DIR, 'Payroll_Module_Task_Completion_Report.docx')
    doc.save(out_path)
    print(f'Document saved: {out_path}')
    print(f'Total images: {added}')

if __name__ == '__main__':
    main()
